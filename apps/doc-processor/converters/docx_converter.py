from docx import Document

def extract_text(file_path: str) -> str:
    """Извлекает текст из DOCX файла"""
    try:
        doc = Document(file_path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
        
        return "\n\n".join(text)
    except Exception as e:
        raise Exception(f"Failed to extract DOCX text: {str(e)}")


